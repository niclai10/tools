import os
from docx import Document
import win32com.client as win32
import tkinter as tk
from tkinter import filedialog, messagebox


def convert_doc_to_docx_wps(doc_path):
    """
    使用 WPS 的 COM 接口将 .doc 文件转换为 .docx 文件
    :param doc_path: .doc 文件的路径
    :return: 转换后的 .docx 文件的路径
    """
    wps = None
    doc = None
    try:
        # 尝试连接 WPS 的 COM 接口
        wps = win32.gencache.EnsureDispatch('KWPS.Application')
        doc = wps.Documents.Open(doc_path)
        docx_path = os.path.splitext(doc_path)[0] + '.docx'
        doc.SaveAs(docx_path, FileFormat=16)  # 16 表示 .docx 格式
        return docx_path
    except Exception as e:
        print(f"转换文件 {doc_path} 时出错: {e}")
    finally:
        if doc:
            doc.Close()
        if wps:
            wps.Quit()


def merge_word_files(input_folder, output_folder):
    """
    合并指定文件夹中的所有 Word 文件（.doc 和 .docx）到一个新的 Word 文件中
    :param input_folder: 包含 Word 文件的文件夹路径
    :param output_folder: 合并后输出文件的文件夹路径
    """
    # 创建一个新的 Word 文档
    merged_document = Document()

    # 获取指定文件夹中的所有 Word 文件，过滤掉临时文件
    word_files = []
    for root, dirs, files in os.walk(input_folder):
        for file in files:
            if file.endswith(('.doc', '.docx')) and not file.startswith('~$'):
                word_files.append(os.path.join(root, file))

    # 按文件名排序
    word_files.sort()

    # 遍历每个 Word 文件
    for file in word_files:
        if file.endswith('.doc'):
            # 将 .doc 文件转换为 .docx 文件
            docx_file = convert_doc_to_docx_wps(file)
            if docx_file:
                doc = Document(docx_file)
                # 删除临时转换的 .docx 文件
                os.remove(docx_file)
            else:
                continue
        else:
            doc = Document(file)

        # 将当前文档的每个段落添加到合并后的文档中
        for paragraph in doc.paragraphs:
            merged_document.add_paragraph(paragraph.text)

        # 将当前文档的每个表格添加到合并后的文档中
        for table in doc.tables:
            new_table = merged_document.add_table(rows=len(table.rows), cols=len(table.columns))
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.cell(i, j).text = cell.text

    # 构建输出文件的完整路径
    output_file = os.path.join(output_folder, "merged_document.docx")
    # 保存合并后的文档
    merged_document.save(output_file)
    messagebox.showinfo("完成", f"合并完成，输出文件: {output_file}")


def select_input_folder():
    folder = filedialog.askdirectory()
    input_folder_entry.delete(0, tk.END)
    input_folder_entry.insert(0, folder)


def select_output_folder():
    folder = filedialog.askdirectory()
    output_folder_entry.delete(0, tk.END)
    output_folder_entry.insert(0, folder)


def start_merge():
    input_folder = input_folder_entry.get()
    output_folder = output_folder_entry.get()

    if not os.path.exists(input_folder):
        messagebox.showerror("错误", f"输入目录 {input_folder} 不存在，请检查路径。")
        return
    if not os.path.exists(output_folder):
        try:
            os.makedirs(output_folder)
        except Exception as e:
            messagebox.showerror("错误", f"创建输出目录 {output_folder} 时出错: {e}")
            return

    merge_word_files(input_folder, output_folder)


# 创建主窗口
root = tk.Tk()
root.title("Word 文件合并工具")

# 输入待合并文件目录
input_folder_label = tk.Label(root, text="待合并文件所在目录:")
input_folder_label.pack(pady=5)
input_folder_entry = tk.Entry(root, width=50)
input_folder_entry.pack(pady=5)
input_folder_button = tk.Button(root, text="选择目录", command=select_input_folder)
input_folder_button.pack(pady=5)

# 输入合并后文件输出目录
output_folder_label = tk.Label(root, text="合并后文件输出目录:")
output_folder_label.pack(pady=5)
output_folder_entry = tk.Entry(root, width=50)
output_folder_entry.pack(pady=5)
output_folder_button = tk.Button(root, text="选择目录", command=select_output_folder)
output_folder_button.pack(pady=5)

# 开始合并按钮
start_button = tk.Button(root, text="开始合并", command=start_merge)
start_button.pack(pady=20)

# 运行主循环
root.mainloop()
